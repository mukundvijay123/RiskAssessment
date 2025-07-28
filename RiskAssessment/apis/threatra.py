from fastapi import APIRouter, Depends, Query,HTTPException,Path
from fastapi.responses import StreamingResponse
import requests
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import distinct
from .db import get_db
from typing import List,Dict
from .tables import ThreatRisk
from .models import ThreatRiskGenerationRequest,ThreatRiskGenerationResponse,ThreatRiskModel,ThreatRiskUpdate
from io import BytesIO
from openpyxl import Workbook
from docx import Document

threatrouter = APIRouter(prefix="/threatRiskAssessment")


def get_threat_risks(orgId: str, db: Session):
    query = db.query(ThreatRisk)
    if orgId != "*":
        query = query.filter(ThreatRisk.organization_id == orgId)
    return query.all()

@threatrouter.post("/generateThreatsRisks")
def generate_threat_risks(request: ThreatRiskGenerationRequest, db: Session = Depends(get_db)):
    payload = {
        "domain": request.domain,
        "category": request.category,
        "business_context": request.business_context or "",
        "specific_focus": request.specific_focus or "",
        "number_of_records": request.number_of_records or 10
    }
    try:
        response=requests.post(
                "https://ey-catalyst-rvce-ey-catalyst.hf.space/threat/api/threat-ra/generate-threat-risks",  # Replace with actual host/port
                json=payload,
                timeout=30  # optional
        )
        print(response.status_code)

        if response.status_code != 200:
            raise HTTPException(status_code=502, detail="LLM API error")

        data = response.json()

        if not data.get("success") or "threatRisks" not in data:
            raise HTTPException(status_code=500, detail="Invalid response from LLM API")
        
        saved_risks = []
        for risk in data["threatRisks"]:
            db_obj = ThreatRisk(
                organization_id=request.organization_id,
                domain=risk["domain"],
                riskName=risk["riskName"],
                threat=risk["threat"],
                vulnerability=risk["vulnerability"],
                category=risk["category"],
                likelihood=risk["likelihood"],
                impact=risk["impact"],
                rating=risk["rating"],
                likelihood_justification=risk["likelihood_justification"],
                impact_justification=risk["impact_justification"],
                threat_justification=risk["threat_justification"],
                vulnerability_justification=risk["vulnerability_justification"]
            )
            db.add(db_obj)
            saved_risks.append(db_obj)

        pydantic_risks = [ThreatRiskModel.from_orm(obj) for obj in saved_risks]

        db.commit()
        return ThreatRiskGenerationResponse(
                success=True,
                threatRisks=saved_risks,
                message=data.get("message", "Records saved")
        )
    except requests.RequestException as e:
        raise HTTPException(status_code=500, detail=f"Failed to reach LLM API: {str(e)}")
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Internal error: {str(e)}")
    


@threatrouter.get("/getRisks", response_model=List[ThreatRiskModel])
def get_risks(orgId: str = Query(...), db: Session = Depends(get_db)):
    if orgId == "*":
        risks = db.query(ThreatRisk).all()
    else:
        risks = db.query(ThreatRisk).filter(ThreatRisk.organization_id == orgId).all()

    return risks



@threatrouter.put("/api/threat-risks/{id}")
def update_threat_risk(
    id: str = Path(...),
    updated_data: ThreatRiskUpdate = ...,
    db: Session = Depends(get_db)
):
    # Fetch the existing record
    threat_risk = db.query(ThreatRisk).filter(ThreatRisk.id == id).first()
    if not threat_risk:
        raise HTTPException(status_code=404, detail="Threat Risk not found")

    # Update fields
    for key, value in updated_data.dict().items():
        setattr(threat_risk, key, value)

    db.commit()
    db.refresh(threat_risk)

    return {
        "success": True,
        "threatRisk": {
            "id": threat_risk.id,
            "domain": threat_risk.domain,
            "riskName": threat_risk.riskName,
            "threat": threat_risk.threat,
            "vulnerability": threat_risk.vulnerability,
            "category": threat_risk.category,
            "likelihood": threat_risk.likelihood,
            "impact": threat_risk.impact,
            "rating": threat_risk.rating,
        }
    }

@threatrouter.get("/domains", response_model=list[str])
def get_domains(orgId: str = Query("*"), db: Session = Depends(get_db)):
    query = db.query(distinct(ThreatRisk.domain))
    if orgId != "*":
        query = query.filter(ThreatRisk.organization_id == orgId)
    return [row[0] for row in query.all()]



@threatrouter.get("/risk-names", response_model=list[str])
def get_risk_names(orgId: str = Query("*"), db: Session = Depends(get_db)):
    query = db.query(distinct(ThreatRisk.riskName))
    if orgId != "*":
        query = query.filter(ThreatRisk.organization_id == orgId)
    return [row[0] for row in query.all()]



@threatrouter.get("/threats", response_model=list[str])
def get_threats(orgId: str = Query("*"), db: Session = Depends(get_db)):
    query = db.query(distinct(ThreatRisk.threat))
    if orgId != "*":
        query = query.filter(ThreatRisk.organization_id == orgId)
    return [row[0] for row in query.all()]



@threatrouter.get("/vulnerabilities", response_model=list[str])
def get_vulnerabilities(orgId: str = Query("*"), db: Session = Depends(get_db)):
    query = db.query(distinct(ThreatRisk.vulnerability))
    if orgId != "*":
        query = query.filter(ThreatRisk.organization_id == orgId)
    return [row[0] for row in query.all()]



@threatrouter.get("/categories", response_model=list[str])
def get_categories(orgId: str = Query("*"), db: Session = Depends(get_db)):
    query = db.query(distinct(ThreatRisk.category))
    if orgId != "*":
        query = query.filter(ThreatRisk.organization_id == orgId)
    return [row[0] for row in query.all()]


@threatrouter.get("/word")
def export_word(orgId: str = Query("*"), db: Session = Depends(get_db)):
    data = get_threat_risks(orgId, db)
    doc = Document()
    doc.add_heading("Threat Risk Records", level=1)

    for risk in data:
        doc.add_heading(risk.riskName, level=2)
        doc.add_paragraph(f"Domain: {risk.domain}")
        doc.add_paragraph(f"Category: {risk.category}")
        doc.add_paragraph(f"Threat: {risk.threat}")
        doc.add_paragraph(f"Vulnerability: {risk.vulnerability}")
        doc.add_paragraph(f"Likelihood: {risk.likelihood}")
        doc.add_paragraph(f"Impact: {risk.impact}")
        doc.add_paragraph(f"Rating: {risk.rating}")
        doc.add_paragraph(f"Likelihood Justification: {risk.likelihood_justification}")
        doc.add_paragraph(f"Impact Justification: {risk.impact_justification}")
        doc.add_paragraph(f"Threat Justification: {risk.threat_justification}")
        doc.add_paragraph(f"Vulnerability Justification: {risk.vulnerability_justification}")
        doc.add_paragraph("")

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=threat_risks.docx"}
    )